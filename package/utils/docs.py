import re
from bs4 import BeautifulSoup
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown import markdown
from docx import Document
from docx.shared import Pt, RGBColor, Inches

tnr = "Times New Roman"
size = Pt(12)
color = RGBColor(0, 0, 0)


def run(el, line):
    heading_run = el.add_run(line)
    heading_run.font.name = tnr
    heading_run.font.size = size
    heading_run.font.color.rgb = color


def process_strong_tags(text, paragraph):
    while "<strong>" in text and "</strong>" in text:
        before_strong = text.split("<strong>")[0]
        inside_strong = text.split("<strong>")[1].split("</strong>")[0]
        after_strong = text.split("</strong>")[1]

        # Add any non-strong text first
        if before_strong:
            paragraph.add_run(before_strong)

        # Add strong (bold) text
        run_s = paragraph.add_run(inside_strong)
        run_s.bold = True

        # Update the text to process the rest
        text = after_strong

    # If any text is left without <strong> tags, add it
    if text:
        paragraph.add_run(text)


def convert_md_to_docx(content, output_file):
    html = md_to_html(content)

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = tnr
    font.size = size
    font.color.rgb = color

    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.first_line_indent = Inches(0.5)

    # for i in range(1, 9):
    #     hsfont = doc.styles[f"Heading {i}"].font
    #     hsfont.name = "Times New Roman"

    for line in html.splitlines():
        if line.startswith("<h1>"):
            heading = doc.add_heading(level=1)
            run(heading, line[4:-5])
            # heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("<h2>"):
            heading = doc.add_heading(level=2)
            run(heading, line[4:-5])
        elif line.startswith("<h3>"):
            heading = doc.add_heading(line[4:-5], level=3)
            run(heading, line[4:-5])
        elif line.startswith("<h4>"):
            heading = doc.add_heading(line[4:-5], level=4)
            run(heading, line[4:-5])
        elif line.startswith("<h5>"):
            heading = doc.add_heading(line[4:-5], level=5)
            run(heading, line[4:-5])
        elif line.startswith("<p>"):
            # paragraph = doc.add_paragraph(line[3:-4])
            # paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Extract and handle strong tags in the paragraph
            text = line[3:-4]  # Remove <p> and </p>
            paragraph = doc.add_paragraph()

            process_strong_tags(text, paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif line.startswith("<ul>"):
            # Handle unordered list
            continue
        elif line.startswith("<li>"):
            text = line[4:-5]
            paragraph = doc.add_paragraph(style=doc.styles["List Bullet"])

            process_strong_tags(text, paragraph)
        # elif "<strong>" in line and "</strong>" in line:
        #     # Process <strong> text and make it bold
        #     text = line.replace("<strong>", "").replace("</strong>", "")
        #     paragraph = doc.add_paragraph()
        #     run_s = paragraph.add_run(text)
        #     run_s.bold = True  # Apply bold styling

    doc.save(output_file)


def text_to_docx(content, output_file):
    doc = Document()

    doc.add_paragraph(content)

    doc.save(output_file)


def markdown_to_text(markdown_string):
    """ Converts a markdown string to plaintext """

    html = md_to_html(markdown_string)

    # extract text
    soup = BeautifulSoup(html, "html.parser")

    text = "".join(soup.findAll(text=True))

    return text


def md_to_html(markdown_string):
    # md -> html -> text since BeautifulSoup can extract text cleanly
    html = markdown(markdown_string)

    # remove code snippets
    html = re.sub(r"<pre>(.*?)</pre>", " ", html)
    html = re.sub(r"<code>(.*?)</code >", " ", html)

    return html